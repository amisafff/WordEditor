import os
import asyncio
from docx import Document


async def process_template_async(template_path, dat_all, output_file_path):
    doc = Document(template_path)

    data = {}
    total_students = len(dat_all)  # Общее количество студентов

    # Статистика по формам обучения
    dnevna_count = 0
    distanc_count = 0
    dnevna_sum = 0
    distanc_sum = 0
    for i in range(1, 21):  # Предположим, что максимум 20 студентов
        data[f"ФИО{i}"] = ""
        data[f"ФОРМ{i}"] = ""
        data[f"НОМ{i}"] = ""
        data[f"ОЦ{i}"] = ""
        data[f"ОТЛ{i}"] = ""
        data[f"Члены_ГЭК{i}_1"] = ""
        data[f"Члены_ГЭК{i}_2"] = ""
        data[f"Члены_ГЭК{i}_3"] = ""
        data[f"Председатель_ГЭК"] = ""
        data[f"Консультант{i}"] = ""
        data[f"Руководитель{i}"] = ""
        data[f"ТЕМА{i}"] = ""
        data[f"ВИЗА{i}"] = ""
        data[f"Степень{i}"] = ""
        data[f"День"] = ""
        data[f"Год"] = ""
        data[f"Время"] = ""  # Можно задать фиксированное время, если нужно
        data[f"ОТЛ{i}"] = ""  # Диплом с отличием

    # Заполнение данных для студентов
    for index, record in enumerate(dat_all):
        # Получаем строку с членами ГЭК
        geks_str = record[6]  # Столбец с членами ГЭК
        geks = geks_str.split(',')  # Если члены разделены запятой
        geks = [geek.strip() for geek in geks]

        substringsDate = record[1].split('-')
        # Сохраняем каждого члена ГЭК
        for g_index, member in enumerate(geks):
            data[f"Члены_ГЭК{index + 1}_{g_index + 1}"] = member
        data[f"ФИО{index + 1}"] = record[2]  # Столбец "ФИО"
        data[f"Специальность"] = record[3]  # Столбец "Специальность"
        data[f"ТЕМА{index + 1}"] = record[4]  # Столбец "Тема ДП"
        data[f"Председатель_ГЭК"] = record[5]  # Столбец "Председатель ГЭК"
        data[f"Консультант{index + 1}"] = record[7]  # Столбец "Консультант"
        data[f"ФОРМ{index + 1}"] = record[8]  # Столбец "Форма обучения"
        data[f"Руководитель{index + 1}"] = record[9]  # Столбец "Руководитель"
        data[f"ОЦ{index + 1}"] = record[10]  # Столбец "Оценка"
        data[f"ВИЗА"] = record[11]  # Столбец "Виза лица, составившего протокол"
        data[f"Степень{index + 1}"] = record[12]  # Столбец "Степень"
        data[f"НОМ{index + 1}"] = record[0]  # Столбец "Номер протокола"
        data[f"День"] = substringsDate[2]
        data[f"Год"] = substringsDate[0]
        data[f"Время"] = substringsDate[3]
        # Логика для "Диплом с отличием"
        if record[13] == "Да":
            data[f"ОТЛ{index + 1}"] = "с отличием"
        else:
            data[f"ОТЛ{index + 1}"] = ""  # Пустая строка, если не с отличием

        # Статистика по форме обучения
        form = record[8]
        grade = int(record[10])

        if form == "Дневная":
            dnevna_count += 1
            dnevna_sum += grade
        elif form == "Дистанционная":
            distanc_count += 1
            distanc_sum += grade

    # Средний балл для каждой формы обучения
    if dnevna_count > 0:
        dnevna_avg = dnevna_sum / dnevna_count
    else:
        dnevna_avg = 0

    if distanc_count > 0:
        distanc_avg = distanc_sum / distanc_count
    else:
        distanc_avg = 0

    # Добавляем статистику в данные
    data["Дневная_СРЕДНИЙ_БАЛЛ"] = round(dnevna_avg, 2)
    data["Дистанционная_СРЕДНИЙ_БАЛЛ"] = round(distanc_avg, 2)

    # Распределение по оценкам для дневной и дистанционной форм
    for grade in range(10, 2, -1):  # Для оценок с 10 по 3
        data[f"Дневная_ОЦ_{grade}_КОЛ"] = sum(
            1 for record in dat_all if record[8] == "Дневная" and int(record[10]) == grade)
        data[f"Дневная_ОЦ_{grade}_ПРЦ"] = round(data[f"Дневная_ОЦ_{grade}_КОЛ"] / dnevna_count * 100,
                                                2) if dnevna_count > 0 else 0

        data[f"Дистанционная_ОЦ_{grade}_КОЛ"] = sum(
            1 for record in dat_all if record[8] == "Дистанционная" and int(record[10]) == grade)
        data[f"Дистанционная_ОЦ_{grade}_ПРЦ"] = round(data[f"Дистанционная_ОЦ_{grade}_КОЛ"] / distanc_count * 100,
                                                      2) if distanc_count > 0 else 0

    # Заменяем метки в параграфах
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

    # Заменяем метки в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))

    # Сохраняем файл
    file_name = os.path.basename(template_path)
    save_path = os.path.join(output_file_path, file_name)
    doc.save(save_path)
    print(f"Файл {file_name} успешно сохранен в {output_file_path}.")


async def fill_word_template_async(dat_all, output_file_path):
    if not os.path.exists(output_file_path):
        os.makedirs(output_file_path)

    template_paths = [
        'DocFile/График защиты ДП _на дверь_для КП.docx',
        'DocFile/График защиты ДП _с оценками_15.06.24_ для КП.docx',
        'DocFile/Присвоение квалификации_№1_15.06.24_для КП.docx',
        'DocFile/Списки студентов, тем, оценки рецензентов_для членов ГЭК_ для КП.docx',
        'DocFile/Статистика ГЭК_для КП.docx'
    ]

    tasks = []
    for template_path in template_paths:
        tasks.append(process_template_async(template_path, dat_all, output_file_path))

    # Выполнение всех асинхронных задач
    await asyncio.gather(*tasks)


# Основная асинхронная функция
def run_async(dat_all, output_file_path):
    asyncio.run(fill_word_template_async(dat_all, output_file_path))
