import os
from docx import Document


def fill_word_template(dat_all, template_path, output_file_path):
    print(dat_all)
    for i in range(len(dat_all)):

        doc = Document(template_path)

        substrings = dat_all[i][6].split(',')
        if len(substrings) <= 6:
            ind = len(substrings)
            for j in range(ind, 6):
                substrings.append("")

        substringsDate = dat_all[i][1].split('-')

        data = {
            "Номер_протокола": dat_all[i][0],
            "День": substringsDate[2],
            "Год": substringsDate[0],
            "ФИО": dat_all[i][2],
            "Специальность": dat_all[i][3],
            "Тема_ДП": dat_all[i][4],
            "Председатель_ГЭК": dat_all[i][5],
            "Члены_ГЭК1": substrings[0],
            "Члены_ГЭК2": substrings[1],
            "Члены_ГЭК3": substrings[2],
            "Члены_ГЭК4": substrings[3],
            "Члены_ГЭК5": substrings[4],
            "Члены_ГЭК6": substrings[5],
            "Руководитель": dat_all[i][9],
            "Консультант": dat_all[i][7],
            "Виза лица": dat_all[i][11],
        }

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{{{key}}}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{{{key}}}}}", str(value))

        file_name = f"{dat_all[i][2]}_.docx"
        save_path = os.path.join(output_file_path, file_name)

        doc.save(save_path)
